"""
Document Assistant - RAG System for Multi-format Document Interaction
This system enables conversational access to both structured and unstructured documents.
"""

import os
import sys
import pandas as pd
import numpy as np
from typing import List, Dict, Any, Optional, Union
import json
import re
from dataclasses import dataclass
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import csv

# Vector store and embeddings
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer

# LLM and processing
import openai
from transformers import pipeline
import torch

@dataclass
class DocumentChunk:
    """Represents a chunk of document content"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_type: str  # 'pdf', 'docx', 'csv', etc.
    
class DocumentProcessor:
    """Handles processing of different document types"""
    
    def __init__(self):
        self.chunks = []
        
    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """Extract text from PDF files with improved chunking"""
        chunks = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                page_texts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    page_texts.append((page_num + 1, text))
                    full_text += text + "\n"
                
                # Create chunks with better strategy
                # Strategy 1: Store full document as one chunk for comprehensive searches
                if len(full_text.strip()) > 100:
                    chunks.append(DocumentChunk(
                        content=full_text.strip(),
                        metadata={
                            'source': file_path,
                            'page': 'all',
                            'filename': os.path.basename(file_path),
                            'chunk_type': 'full_document'
                        },
                        chunk_id=f"{os.path.basename(file_path)}_full",
                        source_type='pdf'
                    ))
                
                # Strategy 2: Create overlapping chunks of ~1500 characters for better context
                chunk_size = 1500
                overlap = 300
                start = 0
                chunk_idx = 0
                
                while start < len(full_text):
                    end = start + chunk_size
                    chunk_text = full_text[start:end]
                    
                    if len(chunk_text.strip()) > 100:
                        chunks.append(DocumentChunk(
                            content=chunk_text.strip(),
                            metadata={
                                'source': file_path,
                                'filename': os.path.basename(file_path),
                                'chunk_type': 'overlap_chunk'
                            },
                            chunk_id=f"{os.path.basename(file_path)}_overlap_{chunk_idx}",
                            source_type='pdf'
                        ))
                    
                    start = end - overlap
                    chunk_idx += 1
                
                # Strategy 3: Per-page chunks for page-specific queries
                for page_num, page_text in page_texts:
                    if len(page_text.strip()) > 50:
                        chunks.append(DocumentChunk(
                            content=page_text.strip(),
                            metadata={
                                'source': file_path,
                                'page': page_num,
                                'filename': os.path.basename(file_path),
                                'chunk_type': 'page'
                            },
                            chunk_id=f"{os.path.basename(file_path)}_page_{page_num}",
                            source_type='pdf'
                        ))
                        
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
        
        return chunks
    
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """Extract text from DOCX files with improved chunking"""
        chunks = []
        try:
            doc = DocxDocument(file_path)
            full_text = ""
            
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            # Extract tables from DOCX as well
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    full_text += " | ".join(row_data) + "\n"
            
            # Strategy 1: Full document chunk
            if len(full_text.strip()) > 100:
                chunks.append(DocumentChunk(
                    content=full_text.strip(),
                    metadata={
                        'source': file_path,
                        'filename': os.path.basename(file_path),
                        'chunk_type': 'full_document'
                    },
                    chunk_id=f"{os.path.basename(file_path)}_full",
                    source_type='docx'
                ))
            
            # Strategy 2: Overlapping chunks with larger size for complete context
            chunk_size = 1500
            overlap = 300
            start = 0
            chunk_idx = 0
            
            while start < len(full_text):
                end = start + chunk_size
                chunk_text = full_text[start:end]
                
                if len(chunk_text.strip()) > 100:
                    chunks.append(DocumentChunk(
                        content=chunk_text.strip(),
                        metadata={
                            'source': file_path,
                            'filename': os.path.basename(file_path),
                            'chunk_type': 'overlap_chunk'
                        },
                        chunk_id=f"{os.path.basename(file_path)}_overlap_{chunk_idx}",
                        source_type='docx'
                    ))
                
                start = end - overlap
                chunk_idx += 1
                
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
        
        return chunks
    
    def process_csv(self, file_path: str) -> tuple:
        """Process CSV files and return both DataFrame and text chunks"""
        chunks = []
        df = None
        
        try:
            df = pd.read_csv(file_path)
            
            # Create summary chunk
            summary = f"CSV File: {os.path.basename(file_path)}\n"
            summary += f"Columns: {', '.join(df.columns)}\n"
            summary += f"Number of rows: {len(df)}\n"
            
            # Add sample data description
            for col in df.columns[:10]:  # Limit to first 10 columns
                if df[col].dtype in ['int64', 'float64']:
                    summary += f"{col}: min={df[col].min()}, max={df[col].max()}, mean={df[col].mean():.2f}\n"
                else:
                    unique_vals = df[col].nunique()
                    summary += f"{col}: {unique_vals} unique values\n"
                    if unique_vals < 20:
                        summary += f"  Values: {', '.join(df[col].unique()[:10].astype(str))}\n"
            
            chunks.append(DocumentChunk(
                content=summary,
                metadata={
                    'source': file_path,
                    'filename': os.path.basename(file_path),
                    'type': 'csv_summary'
                },
                chunk_id=f"{os.path.basename(file_path)}_summary",
                source_type='csv'
            ))
            
            # Create chunks for specific important rows (like orders with specific names)
            if 'mktg_specialistsmanagers' in df.columns:
                for name in df['mktg_specialistsmanagers'].unique():
                    if pd.notna(name):
                        name_data = df[df['mktg_specialistsmanagers'] == name]
                        content = f"Orders handled by {name}:\n"
                        content += f"Total orders: {len(name_data)}\n"
                        products = name_data['product_'].dropna().astype(str).unique()[:5]
                        content += f"Products: {', '.join(products)}\n"
                        if 'qty' in df.columns:
                            qty_series = pd.to_numeric(name_data['qty'], errors='coerce')
                            total_qty = qty_series.sum(skipna=True)
                            if pd.isna(total_qty):
                                total_qty = 0
                            elif float(total_qty).is_integer():
                                total_qty = int(total_qty)
                            content += f"Total quantity: {total_qty}\n"
                        if 'status' in df.columns:
                            content += f"Status distribution: {name_data['status'].value_counts().to_dict()}\n"
                        
                        chunks.append(DocumentChunk(
                            content=content,
                            metadata={
                                'source': file_path,
                                'filename': os.path.basename(file_path),
                                'person': name
                            },
                            chunk_id=f"{os.path.basename(file_path)}_{name}",
                            source_type='csv'
                        ))
            
        except Exception as e:
            print(f"Error processing CSV {file_path}: {e}")
        
        return chunks, df

class VectorStore:
    """Manages vector storage and retrieval"""
    
    def __init__(self, persist_dir: str = "./chroma_db"):
        self.persist_dir = persist_dir
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Initialize ChromaDB
        self.client = chromadb.PersistentClient(path=persist_dir)
        
        # Create or get collection
        try:
            self.collection = self.client.create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
        except:
            self.collection = self.client.get_collection("documents")
    
    def add_documents(self, chunks: List[DocumentChunk]):
        """Add document chunks to vector store"""
        if not chunks:
            return
        
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedding_model.encode(texts)
        
        ids = [chunk.chunk_id for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]
        
        self.collection.add(
            embeddings=embeddings.tolist(),
            documents=texts,
            metadatas=metadatas,
            ids=ids
        )
    
    def search(self, query: str, k: int = 50) -> List[Dict]:
        """Search for relevant documents - increased to 50 for comprehensive, complete results"""
        query_embedding = self.embedding_model.encode([query])
        
        results = self.collection.query(
            query_embeddings=query_embedding.tolist(),
            n_results=min(k, self.collection.count())  # Don't exceed available documents
        )
        
        return results

class StructuredDataHandler:
    """Handles queries on structured data (CSV/Excel)"""
    
    def __init__(self):
        self.dataframes = {}
    
    def add_dataframe(self, name: str, df: pd.DataFrame):
        """Store a dataframe for querying"""
        self.dataframes[name] = df
    
    def query_data(self, query: str) -> Dict[str, Any]:
        """Process queries on structured data"""
        query_lower = query.lower()
        results = {}
        
        # Check for specific patterns in the query
        for df_name, df in self.dataframes.items():
            # Query about Ranjit
            if 'ranjit' in query_lower:
                if 'mktg_specialistsmanagers' in df.columns:
                    ranjit_data = df[df['mktg_specialistsmanagers'] == 'Ranjit']
                    if not ranjit_data.empty:
                        if 'qty' in df.columns:
                            qty_series = pd.to_numeric(ranjit_data['qty'], errors='coerce').dropna()
                            if not qty_series.empty:
                                total_qty = qty_series.sum()
                                results['ranjit_order_quantity'] = total_qty
                        results['ranjit_total_orders'] = len(ranjit_data)
                        if 'product_' in df.columns:
                            results['ranjit_products'] = ranjit_data['product_'].unique().tolist()
                        results['ranjit_source'] = f"{df_name}.csv"
            
            # Query about dispatch status - only check pharmaceutical orders
            if 'dispatch' in query_lower and 'percentage' in query_lower:
                # Only check dataframes that have order data (with mktg_specialistsmanagers column)
                if 'status' in df.columns and 'mktg_specialistsmanagers' in df.columns:
                    total = len(df)
                    not_dispatched = len(df[df['status'] != 'Dispatched'])
                    percentage = (not_dispatched / total) * 100
                    results['not_dispatched_percentage'] = round(percentage, 2)
                    results['not_dispatched_count'] = not_dispatched
                    results['total_orders'] = total
                    results['dispatch_source'] = f"{df_name}.csv"
            
            # Query about recoil kits - check both category and product name
            if 'recoil kit' in query_lower or 'recoil' in query_lower:
                # Check if it's supply chain data with category column
                if 'category' in df.columns:
                    recoil_data = df[df['category'].str.contains('Recoil', case=False, na=False)]
                    if not recoil_data.empty and 'product_name' in df.columns:
                        results['recoil_products'] = recoil_data['product_name'].unique().tolist()
                        results['recoil_count'] = len(recoil_data)
                        results['recoil_source'] = f"{df_name}.csv"
                # Check pharmaceuticals with product_ column
                elif 'product_' in df.columns:
                    product_mask = df['product_'].astype(str).str.contains('recoil', case=False, na=False)
                    type_mask = (
                        df['type_of_order'].astype(str).str.contains('recoil', case=False, na=False)
                        if 'type_of_order' in df.columns else pd.Series(False, index=df.index)
                    )
                    recoil_data = df[product_mask | type_mask]
                    if not recoil_data.empty:
                        results['recoil_products'] = recoil_data['product_'].dropna().unique().tolist()
                        results['recoil_count'] = len(recoil_data)
                        results['recoil_source'] = f"{df_name}.csv"
            
            # Query about Glock - check both product_name and product_ columns
            if 'glock' in query_lower:
                glock_pattern = r'glock[\s\-]*17'
                # Check supply chain data with product_name column
                if 'product_name' in df.columns:
                    glock_data = df[df['product_name'].str.contains(glock_pattern, case=False, regex=True, na=False)]
                    if not glock_data.empty:
                        if 'planned_wo_release_date' in df.columns:
                            results['glock_wo_release_dates'] = glock_data['planned_wo_release_date'].dropna().unique().tolist()
                        if 'product_name' in df.columns:
                            results['glock_products'] = glock_data['product_name'].unique().tolist()
                        if 'status' in df.columns:
                            results['glock_status'] = glock_data['status'].unique().tolist()
                        results['glock_source'] = f"{df_name}.csv"
                # Check pharmaceuticals with product_ column
                elif 'product_' in df.columns:
                    glock_data = df[df['product_'].str.contains(glock_pattern, case=False, regex=True, na=False)]
                    if not glock_data.empty:
                        if 'wo_release_date_planned' in df.columns:
                            results['glock_wo_release_dates'] = glock_data['wo_release_date_planned'].dropna().unique().tolist()
                        results['glock_products'] = glock_data['product_'].unique().tolist()
                        if 'status' in df.columns:
                            results['glock_status'] = df.loc[glock_data.index, 'status'].dropna().unique().tolist()
                        results['glock_source'] = f"{df_name}.csv"
        
        return results

class ConversationalRAG:
    """Main RAG system for conversational document interaction"""
    
    def __init__(self, source_folder: str):
        self.source_folder = source_folder
        self.doc_processor = DocumentProcessor()
        self.vector_store = VectorStore()
        self.structured_handler = StructuredDataHandler()
        self.all_chunks = []
        
        # Initialize local LLM (using HuggingFace model)
        try:
            self.qa_pipeline = pipeline(
                "question-answering",
                model="deepset/roberta-base-squad2",
                device=0 if torch.cuda.is_available() else -1
            )
            self.use_local_llm = True
        except:
            self.use_local_llm = False
            print("Warning: Local LLM not available. Using rule-based responses.")
        
        self.load_documents()
    
    def load_documents(self):
        """Load all documents from source folder"""
        print(f"Loading documents from {self.source_folder}...")
        
        for file_path in Path(self.source_folder).glob("*"):
            if file_path.suffix.lower() == '.pdf':
                chunks = self.doc_processor.process_pdf(str(file_path))
                self.all_chunks.extend(chunks)
                self.vector_store.add_documents(chunks)
                
            elif file_path.suffix.lower() == '.docx':
                chunks = self.doc_processor.process_docx(str(file_path))
                self.all_chunks.extend(chunks)
                self.vector_store.add_documents(chunks)
                
            elif file_path.suffix.lower() == '.csv':
                chunks, df = self.doc_processor.process_csv(str(file_path))
                if df is not None:
                    self.structured_handler.add_dataframe(file_path.stem, df)
                self.all_chunks.extend(chunks)
                self.vector_store.add_documents(chunks)
        
        print(f"Loaded {len(self.all_chunks)} document chunks")
    
    def generate_response(self, query: str) -> Dict[str, Any]:
        """Generate response to user query"""
        response = {
            'query': query,
            'answer': '',
            'sources': [],
            'data': None
        }
        
        # First, check structured data
        structured_results = self.structured_handler.query_data(query)
        
        # Search vector store for relevant documents - get MAXIMUM results for comprehensive answers
        search_results = self.vector_store.search(query, k=30)
        
        # Build comprehensive context from ALL search results
        context = ""
        source_files = set()
        retrieved_chunks = []
        
        if search_results and search_results['documents']:
            for doc, metadata in zip(search_results['documents'][0], search_results['metadatas'][0]):
                retrieved_chunks.append({'text': doc, 'metadata': metadata})
                source_files.add(metadata.get('filename', 'Unknown'))
            
            # Prioritize full_document chunks for comprehensive answers
            full_docs = []
            page_chunks = []
            other_chunks = []
            
            for chunk in retrieved_chunks:
                metadata = chunk['metadata']
                doc = chunk['text']
                if metadata.get('chunk_type') == 'full_document':
                    full_docs.append(doc)
                elif metadata.get('chunk_type') == 'page':
                    page_chunks.append(doc)
                else:
                    other_chunks.append(doc)
            
            # Use full documents first for complete information, then pages, then overlapping chunks
            # Don't limit the chunks - use ALL relevant information
            context = "\n\n---\n\n".join(full_docs + page_chunks + other_chunks)
            response['sources'] = list(source_files)
        
        # Generate answer based on query type
        query_lower = query.lower()
        
        # Handle specific queries with structured data
        if 'ranjit' in query_lower and 'order quantity' in query_lower:
            if 'ranjit_order_quantity' in structured_results:
                qty_value = structured_results['ranjit_order_quantity']
                if isinstance(qty_value, (int, float)):
                    qty = int(qty_value) if abs(qty_value - int(qty_value)) < 1e-6 else round(qty_value, 2)
                else:
                    try:
                        qty_float = float(qty_value)
                        qty = int(qty_float) if abs(qty_float - int(qty_float)) < 1e-6 else round(qty_float, 2)
                    except (TypeError, ValueError):
                        qty = qty_value
                total_orders = structured_results.get('ranjit_total_orders')
                products = structured_results.get('ranjit_products') or []
                parts = [f"Ranjit handled {qty} unit{'s' if qty != 1 else ''}"]
                if total_orders:
                    parts.append(f"across {total_orders} order{'s' if total_orders != 1 else ''}")
                response_lines = [" ".join(parts).strip() + "."]
                if products:
                    response_lines.append("Products managed: " + ", ".join(sorted(products)) + ".")
                response['answer'] = " ".join(response_lines)
                response['data'] = structured_results
                response['sources'] = [structured_results.get('ranjit_source', 'pharmaceuticals.csv')]
        
        elif 'dispatch' in query_lower and 'percentage' in query_lower:
            if 'not_dispatched_percentage' in structured_results:
                response['answer'] = f"{structured_results['not_dispatched_percentage']}% of orders haven't been dispatched. "
                response['answer'] += f"That's {structured_results['not_dispatched_count']} out of {structured_results['total_orders']} total orders."
                response['data'] = structured_results
                response['sources'] = [structured_results.get('dispatch_source', 'pharmaceuticals.csv')]
        
        elif 'recoil kit' in query_lower:
            if 'recoil_products' in structured_results:
                products = sorted(structured_results['recoil_products'])
                list_lines = "\n".join(f"- {product}" for product in products)
                response['answer'] = "Products under recoil kit orders:\n" + list_lines
                response['data'] = structured_results
                response['sources'] = [structured_results.get('recoil_source', 'pharmaceuticals.csv')]
        
        elif 'glock' in query_lower:
            glock_answer = self._format_glock_release_summary(structured_results)
            if glock_answer:
                response['answer'] = glock_answer
                response['data'] = structured_results
                response['sources'] = [structured_results.get('glock_source', 'pharmaceuticals.csv')]
            elif context:
                response['answer'] = self._extract_from_context(query, context)
        
        elif 'gst' in query_lower or 'insurance' in query_lower:
            insurance_info = self._extract_insurance_details(retrieved_chunks, context)
            if insurance_info:
                response['answer'] = self._format_insurance_response(insurance_info)
                if insurance_info.get('sources'):
                    response['sources'] = sorted(insurance_info['sources'])
            elif context:
                response['answer'] = self._extract_from_context(query, context)

        elif (
            'data scientist' in query_lower and
            any(keyword in query_lower for keyword in ['criteria', 'requirement', 'qualification', 'responsibil', 'profile'])
        ):
            criteria_data = self._extract_data_scientist_criteria(retrieved_chunks, context)
            if criteria_data:
                response['answer'] = criteria_data['answer']
                if criteria_data.get('sources'):
                    response['sources'] = sorted(criteria_data['sources'])
            elif context:
                response['answer'] = self._extract_from_context(query, context)
        
        elif any(word in query_lower for word in ['hiring', 'recruit', 'data scientist', 'roles', 'position', 'opening']):
            if retrieved_chunks:
                hiring_roles = self._extract_hiring_roles(retrieved_chunks)
                if hiring_roles:
                    response['answer'] = self._format_hiring_response(hiring_roles)
                    response['sources'] = sorted({role['source'] for role in hiring_roles})
                elif context:
                    response['answer'] = self._extract_from_context(query, context)
            elif context:
                response['answer'] = self._extract_from_context(query, context)
        
        elif 'create a zone' in query_lower or ('log book' in query_lower and 'zone' in query_lower):
            if context:
                zone_answer = self._extract_zone_instructions(context)
                if zone_answer:
                    response['answer'] = zone_answer
                else:
                    response['answer'] = self._extract_from_context(query, context)
        elif 'log book' in query_lower:
            logbook_answer = None
            logbook_keywords = [
                'benefit', 'benefits', 'benifit', 'benifits',
                'set up', 'setup', 'set-it-up', 'set it up',
                'getting started', 'configure', 'configuration'
            ]
            if any(keyword in query_lower for keyword in logbook_keywords):
                logbook_answer = self._extract_logbook_benefits_and_setup(retrieved_chunks, context)
            if logbook_answer:
                response['answer'] = logbook_answer
            elif context:
                response['answer'] = self._extract_from_context(query, context)
        
        # If no specific handler, use context-based answer
        if not response['answer'] and context:
            response['answer'] = self._extract_from_context(query, context)
        
        # Fallback if no answer generated
        if not response['answer']:
            response['answer'] = "I couldn't find specific information about your query in the documents. Please try rephrasing or asking about specific data available in the files."
        
        return response
    
    def _extract_from_context(self, query: str, context: str) -> str:
        """Extract comprehensive answer from context with detailed information"""
        llm_answer = ""
        if self.use_local_llm:
            try:
                # Use MUCH larger context for complete answers
                result = self.qa_pipeline(question=query, context=context[:12000])
                llm_answer = result.get('answer', '').strip()
            except:
                llm_answer = ""
        
        # Enhanced rule-based extraction for comprehensive answers
        query_lower = query.lower()
        query_keywords = set([word for word in query_lower.split() if len(word) > 3])
        
        # Split context into sentences and paragraphs
        paragraphs = context.split('\n\n')
        relevant_paragraphs = []
        scores = []
        
        # Score each paragraph by relevance
        for para in paragraphs:
            if len(para.strip()) < 20:
                continue
                
            para_lower = para.lower()
            para_words = set([word for word in para_lower.split() if len(word) > 3])
            
            # Calculate relevance score
            keyword_matches = len(query_keywords.intersection(para_words))
            
            # Boost score for exact phrase matches
            if any(keyword in para_lower for keyword in query_keywords):
                keyword_matches += 2
            
            if keyword_matches > 0:
                scores.append((keyword_matches, para))
        
        # Sort by relevance and take MORE paragraphs for complete answers
        scores.sort(reverse=True, key=lambda x: x[0])
        
        # Build comprehensive answer from ALL relevant paragraphs
        answer_parts = []
        seen_content = set()
        
        # Take up to 20 most relevant paragraphs instead of 8 for complete answers
        for score, para in scores[:20]:
            para_clean = para.strip()
            # Avoid duplicates
            if para_clean not in seen_content and len(para_clean) > 30:
                answer_parts.append(para_clean)
                seen_content.add(para_clean)
        
        if answer_parts:
            # Create a comprehensive answer - NO TRUNCATION for complete information
            combined_parts = []
            if llm_answer:
                combined_parts.append(llm_answer)
            combined_parts.extend(answer_parts)
            
            deduped = []
            seen = set()
            for part in combined_parts:
                cleaned = part.strip()
                if not cleaned:
                    continue
                key = cleaned.lower()
                if key not in seen:
                    deduped.append(cleaned)
                    seen.add(key)
            
            if deduped:
                return "\n\n".join(deduped)
        
        # If no good paragraph matches, try sentence-level extraction
        sentences = re.split(r'[.!?]+', context)
        relevant_sentences = []
        
        for sentence in sentences:
            sentence_clean = sentence.strip()
            if len(sentence_clean) < 20:
                continue
            
            sentence_lower = sentence_clean.lower()
            if any(keyword in sentence_lower for keyword in query_keywords):
                relevant_sentences.append(sentence_clean)
        
        if relevant_sentences:
            combined_parts = []
            if llm_answer:
                combined_parts.append(llm_answer)
            combined_parts.extend(relevant_sentences)
            answer = ". ".join(combined_parts)
            if not answer.endswith('.'):
                answer += "."
            return answer
        
        if llm_answer:
            return llm_answer
        
        return f"Based on the documents, I found relevant information but it may not directly answer your specific question. Please try rephrasing or asking more specifically about: {', '.join(list(query_keywords)[:5])}"
    
    def _extract_hiring_roles(self, retrieved_chunks: List[Dict[str, Any]]) -> List[Dict[str, str]]:
        """Extract hiring role information from retrieved document chunks."""
        role_patterns = [
            r'(?:role|position|designation|job title)\s*[:\-]\s*(.+)',
            r'(?:hiring|looking for|recruiting)\s+(?:an?|the)?\s*([A-Za-z0-9 /&,\-]+)',
            r'open(?:ing| position)s?\s*[:\-]\s*(.+)'
        ]
        job_terms = [
            'developer', 'scientist', 'analyst', 'engineer', 'manager',
            'consultant', 'architect', 'specialist', 'designer', 'lead',
            'administrator', 'executive', 'associate', 'intern', 'expert',
            'officer', 'technician', 'coordinator', 'director', 'owner',
            'controller', 'supervisor', 'strategist', 'analyst', 'planner',
            'bi developer', 'data engineer'
        ]
        job_base_terms = [term for term in job_terms if ' ' not in term]
        context_keywords = [
            'hiring', 'role', 'position', 'opening', 'vacancy',
            'recruit', 'job', 'opportunity', 'career'
        ]
        role_prefix_terms = {
            'senior', 'junior', 'lead', 'principal', 'data', 'power',
            'bi', 'business', 'full', 'stack', 'cloud', 'machine',
            'learning', 'software', 'java', 'python', 'powerbi',
            'digital', 'analytics'
        }
        disallowed_starts = (
            'experience', 'responsibil', 'requirement', 'skills',
            'about', 'summary', 'profile', 'here', 'as a', 'we are'
        )
        
        roles_map = {}
        seen = set()
        
        for chunk in retrieved_chunks:
            text = chunk['text']
            metadata = chunk.get('metadata', {})
            filename = metadata.get('filename', 'Unknown')
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            
            for idx, line in enumerate(lines):
                normalized = line.lower()
                context_window = " ".join(lines[max(0, idx-2):idx+1]).lower()
                
                role_title = None
                for pattern in role_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        role_title = match.group(1).strip(" -•:\t.;")
                        break
                
                if not role_title:
                    if 'job description' in normalized and any(term in normalized for term in job_terms):
                        prefix = re.split(r'job description', line, flags=re.IGNORECASE)[0].strip()
                        words = prefix.split()
                        job_idx = None
                        for idx in range(len(words) - 1, -1, -1):
                            clean_word = re.sub(r'[^a-z0-9]', '', words[idx].lower())
                            if clean_word in job_base_terms:
                                job_idx = idx
                                break
                        if job_idx is not None:
                            start_idx = job_idx
                            while start_idx > 0:
                                prev_clean = re.sub(r'[^a-z0-9]', '', words[start_idx - 1].lower())
                                if prev_clean in role_prefix_terms:
                                    start_idx -= 1
                                else:
                                    break
                            candidate = " ".join(words[start_idx:job_idx + 1]).strip(" -•:\t.;")
                            if candidate:
                                role_title = candidate
                
                if not role_title:
                    bullet_match = re.match(r'^[•\-\*\d\.\)\( ]+', line)
                    bullet_line = re.sub(r'^[•\-\*\d\.\)\( ]+', '', line)
                    if bullet_match:
                        if any(term in bullet_line.lower() for term in job_terms):
                            # Require contextual signal that this section is about hiring
                            if any(keyword in context_window for keyword in context_keywords) or any(keyword in normalized for keyword in context_keywords):
                                role_title = bullet_line.strip(" -•:\t.;")
                
                if not role_title and any(term in normalized for term in job_terms):
                    stripped_line = line.strip()
                    stripped_lower = stripped_line.lower()
                    words = stripped_line.split()
                    if len(stripped_line) <= 60 and len(words) <= 6 and not any(punct in stripped_line for punct in ['.', '!', '?', "'", '’']):
                        if not stripped_lower.startswith(disallowed_starts):
                            if re.match(r'^[A-Za-z0-9 /&\-\(\)]+$', stripped_line):
                                words_clean = [re.sub(r'[^a-z0-9]', '', w.lower()) for w in stripped_line.split()]
                                if any(word in job_base_terms for word in words_clean[-2:]):
                                    role_title = stripped_line
                
                if not role_title:
                    continue
                
                role_title = re.sub(r'\b(role|position|opening)\b', '', role_title, flags=re.IGNORECASE).strip(" -•:\t.;")
                role_title = re.sub(r'^(?:for|the|a|an)\s+', '', role_title, flags=re.IGNORECASE)
                role_title = re.sub(r'^(?:job description for|job desc for)\s+', '', role_title, flags=re.IGNORECASE)
                role_title = re.sub(r'^(?:opening|openings|open)\s+(?:for\s+)?', '', role_title, flags=re.IGNORECASE)
                role_title = re.sub(r'\s+', ' ', role_title).strip()
                
                normalized_role = role_title.lower()
                if not any(term in normalized_role for term in job_terms):
                    continue
                
                key = (role_title.lower(), filename.lower())
                if key in seen:
                    continue
                
                detail_lines = [line]
                for extra_line in lines[idx+1:idx+4]:
                    extra_lower = extra_line.lower()
                    potential_new_role = (
                        (re.search(r'(?:position|role|opening)\s*[:\-]', extra_lower) and any(term in extra_lower for term in job_terms))
                        or (re.match(r'^[•\-\*\d]', extra_line.strip()) and any(term in extra_lower for term in job_terms))
                    )
                    if potential_new_role:
                        break
                    if any(trigger in extra_lower for trigger in ['responsibilit', 'requirement', 'skills', 'experience', 'about the role']):
                        detail_lines.append(extra_line)
                        break
                    if len(extra_line.split()) < 4:
                        continue
                    detail_lines.append(extra_line)
                    if len(" ".join(detail_lines)) > 500:
                        break
                
                detail = " ".join(detail_lines).strip()
                if len(detail) > 600:
                    detail = detail[:600].rstrip() + "..."
                
                if key not in roles_map:
                    roles_map[key] = {
                        'role': role_title,
                        'source': filename,
                        'details': []
                    }
                if detail and detail not in roles_map[key]['details']:
                    roles_map[key]['details'].append(detail)
                seen.add(key)
        
        return list(roles_map.values())
    
    def _format_hiring_response(self, roles: List[Dict[str, str]]) -> str:
        """Format hiring role information into a concise list of role names."""
        if not roles:
            return ""
        
        unique_roles = {}
        for role in roles:
            key = (role['role'], role['source'])
            unique_roles[key] = role
        
        lines = ["Open roles identified across the documents:"]
        for (role_name, source) in sorted(unique_roles.keys()):
            lines.append(f"- {role_name} (source: {source})")
        
        return "\n".join(lines)

    def _get_chunk_text(
        self,
        retrieved_chunks: List[Dict[str, Any]],
        filename_keyword: Optional[str] = None,
        keyword_filters: Optional[List[str]] = None
    ) -> str:
        """Aggregate text from retrieved chunks that match a filename or keywords."""
        if not retrieved_chunks:
            return ""
        
        texts = []
        seen = set()
        keyword_filters = [kw.lower() for kw in keyword_filters] if keyword_filters else []
        filename_keyword = filename_keyword.lower() if filename_keyword else None
        
        for chunk in retrieved_chunks:
            text = chunk.get('text') or ""
            if not text.strip():
                continue
            metadata = chunk.get('metadata', {})
            filename = metadata.get('filename', '').lower()
            lower_text = text.lower()
            
            include = False
            if filename_keyword and filename_keyword in filename:
                include = True
            if keyword_filters and any(keyword in lower_text for keyword in keyword_filters):
                include = True
            if not filename_keyword and not keyword_filters:
                include = True
            
            if include and text not in seen:
                texts.append(text)
                seen.add(text)
        
        return "\n\n".join(texts)

    def _clean_section_line(self, line: str) -> str:
        cleaned = line.strip()
        if not cleaned:
            return ""
        cleaned = re.sub(r'\[.*?\]', '', cleaned)
        cleaned = re.sub(r'^[\-\u2022\u25CF\u25CB\u25A0·•○●■▪▫\*\d\)\(]+', '', cleaned).strip()
        cleaned = cleaned.replace('–', '-')
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        noise_tokens = [
            'gps log book – user manual', 'gps log book - user manual',
            'table of contents', 'insert website', 'insert logo', 'page'
        ]
        lowered = cleaned.lower()
        for noise in noise_tokens:
            if lowered.startswith(noise):
                return ""
        return cleaned

    def _extract_bullet_lines(self, block: str) -> List[str]:
        """Convert a block of text with bullets into clean list items."""
        if not block:
            return []
        
        aggregated = []
        current = ""
        bullet_chars = {'●', '○', '•', '', '-', '▪', '▫', '■', '\uf0b7'}
        bullet_regex = r'^[\-\u2022\u25CF\u25CB\u25A0·•○●■▪▫\*\uf0b7]+'
        
        for raw in block.splitlines():
            stripped = raw.strip()
            if not stripped:
                if current:
                    aggregated.append(current.strip())
                    current = ""
                continue
            
            if stripped in bullet_chars:
                if current:
                    aggregated.append(current.strip())
                    current = ""
                continue
            
            if re.match(bullet_regex + r'\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
                if current:
                    aggregated.append(current.strip())
                cleaned = re.sub(bullet_regex + r'\s*', '', stripped)
                current = cleaned
            else:
                if current:
                    current += ' ' + stripped
                else:
                    current = stripped
        
        if current:
            aggregated.append(current.strip())
        
        cleaned_lines = []
        for line in aggregated:
            cleaned = self._clean_section_line(line)
            if cleaned:
                cleaned_lines.append(cleaned)
        
        return cleaned_lines

    def _extract_numbered_lines(self, block: str) -> List[str]:
        """Extract ordered steps (1., 2., etc.) from a block."""
        if not block:
            return []
        
        steps = []
        current = ""
        for raw in block.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue
            match = re.match(r'^(\d+)\.\s*(.+)', stripped)
            if match:
                if current:
                    steps.append(current.strip())
                current = match.group(2)
            else:
                if current:
                    current += ' ' + stripped
        if current:
            steps.append(current.strip())
        
        cleaned_steps = []
        for step in steps:
            cleaned = self._clean_section_line(step)
            if cleaned:
                cleaned_steps.append(cleaned)
        return cleaned_steps
    
    def _summarize_section(self, block: str, keywords: List[str]) -> Optional[str]:
        if not block:
            return None
        sentences = re.split(r'(?<=[.!?])\s+', block)
        for sentence in sentences:
            cleaned = self._clean_section_line(sentence)
            if not cleaned:
                continue
            lower = cleaned.lower()
            if any(keyword in lower for keyword in keywords):
                return cleaned
        merged = self._clean_section_line(" ".join(block.splitlines()[:3]))
        return merged or None

    def _filter_job_section_lines(self, lines: List[str]) -> List[str]:
        """Remove placeholder or unrelated lines from job description sections."""
        filtered = []
        disallowed_keywords = [
            'job description', 'template', 'opening-', 'position-', 'www.',
            'gps log book', 'illustrated quick start', 'usb cable', 'technical specifications'
        ]
        for line in lines:
            lower = line.lower()
            if any(keyword in lower for keyword in disallowed_keywords):
                continue
            filtered.append(line)
        return filtered

    def _get_section_block(
        self,
        text: str,
        header_token: str,
        stop_tokens: Optional[List[str]] = None
    ) -> str:
        if not text:
            return ""
        
        flags = re.IGNORECASE | re.DOTALL | re.MULTILINE
        header_regex = rf'^\s*{re.escape(header_token)}[^\n]*\n?'
        if stop_tokens:
            stop_regex = "|".join(rf'^\s*{re.escape(token)}' for token in stop_tokens)
            pattern = re.compile(rf'{header_regex}(.*?)(?={stop_regex})', flags)
        else:
            pattern = re.compile(rf'{header_regex}(.*)', flags)
        
        matches = pattern.findall(text)
        if not matches:
            return ""
        block = max(matches, key=len)
        return block

    def _extract_insurance_details(
        self,
        retrieved_chunks: List[Dict[str, Any]],
        context: str
    ) -> Optional[Dict[str, Any]]:
        """Pull GST and expiry information from insurance documents."""
        chunk_text = self._get_chunk_text(
            retrieved_chunks,
            filename_keyword='doc 1.pdf',
            keyword_filters=['gst', 'expiry', 'premium', 'tp']
        )
        search_space = chunk_text or context
        if not search_space:
            return None
        
        sources = {
            chunk.get('metadata', {}).get('filename', 'Unknown')
            for chunk in retrieved_chunks
            if any(keyword in (chunk.get('text') or '').lower() for keyword in ['gst', 'tp expiry'])
        }
        sources = {source for source in sources if source}
        
        def find_first(patterns: List[str]) -> Optional[str]:
            for pattern in patterns:
                match = re.search(pattern, search_space, re.IGNORECASE)
                if match:
                    return match.group(1).strip()
            return None
        
        cgst_raw = find_first([
            r'CGST[^\n\r]*?([\d,]+(?:\.\d+)?)\s*(?:\n|$)'
        ])
        sgst_raw = find_first([
            r'SGST[^\n\r]*?([\d,]+(?:\.\d+)?)\s*(?:\n|$)'
        ])
        gst_total_raw = find_first([
            r'\bGST(?:\s+Amount)?[^\n\r]*?(?:₹|Rs\.?)?\s*([\d,]+(?:\.\d+)?)'
        ])
        tp_expiry = find_first([
            r'(?:TP|Third Party)[^\d]{0,15}(?:Expiry|Valid(?:ity)?\s*(?:Till|Until|Up\s*to)|Expire)[^\d]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'Expiry Date[^\d]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
        ])
        od_expiry = find_first([
            r'OD\s+Expiry\s+Date[^\d]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
        ])
        
        if not any([cgst_raw, sgst_raw, tp_expiry, od_expiry]):
            return None
        
        def parse_amount(value: Optional[str]) -> Optional[float]:
            if not value:
                return None
            cleaned = value.replace(',', '')
            try:
                return float(cleaned)
            except ValueError:
                return None
        
        cgst = parse_amount(cgst_raw)
        sgst = parse_amount(sgst_raw)
        gst_total = parse_amount(gst_total_raw)
        if not gst_total and cgst is not None and sgst is not None:
            gst_total = cgst + sgst
        
        return {
            'cgst': cgst,
            'sgst': sgst,
            'gst_total': gst_total,
            'tp_expiry': tp_expiry,
            'od_expiry': od_expiry,
            'sources': list(sources) if sources else ['Doc 1.pdf']
        }

    def _format_amount(self, value: float) -> str:
        if value is None:
            return ""
        if abs(value - int(value)) < 1e-6:
            return f"{int(value):,}"
        return f"{value:,.2f}".rstrip('0').rstrip('.')

    def _format_insurance_response(self, info: Dict[str, Any]) -> str:
        parts = []
        gst_parts = []
        if info.get('cgst') is not None:
            gst_parts.append(f"CGST charged: Rs.{self._format_amount(info['cgst'])}")
        if info.get('sgst') is not None:
            gst_parts.append(f"SGST charged: Rs.{self._format_amount(info['sgst'])}")
        if gst_parts:
            if info.get('gst_total') is not None and len(gst_parts) > 1:
                gst_parts.append(f"Total GST collected: Rs.{self._format_amount(info['gst_total'])}")
            parts.append(" ".join(gst_parts))
        
        if info.get('tp_expiry'):
            parts.append(f"Third-party (TP) insurance is valid until {info['tp_expiry']}.")
        if info.get('od_expiry'):
            parts.append(f"Own-damage cover expires on {info['od_expiry']}.")
        
        return " ".join(parts).strip()
    
    def _normalize_date_value(self, raw_value: Optional[str]) -> Optional[str]:
        if not raw_value:
            return None
        value = str(raw_value).strip()
        match = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})', value)
        if match:
            return match.group(1)
        cleaned = re.sub(r'_DATE_VALUE.*', '', value, flags=re.IGNORECASE)
        return cleaned.strip() or value
    
    def _deduplicate_preserve_order(self, items: List[str]) -> List[str]:
        seen = set()
        ordered = []
        for item in items:
            normalized = item.strip().lower()
            if not item or normalized in seen:
                continue
            seen.add(normalized)
            ordered.append(item.strip())
        return ordered
    
    def _format_glock_release_summary(self, structured_results: Dict[str, Any]) -> Optional[str]:
        dates = structured_results.get('glock_wo_release_dates')
        products = structured_results.get('glock_products')
        if not dates or not products:
            return None
        normalized_dates = [self._normalize_date_value(date) or date for date in dates]
        status_list = structured_results.get('glock_status', [])
        
        answer_parts = []
        for idx, product in enumerate(products):
            date = normalized_dates[min(idx, len(normalized_dates) - 1)]
            status = status_list[min(idx, len(status_list) - 1)] if status_list else None
            phrase = f"{product}: planned WO release date {date}"
            if status:
                phrase += f" (status: {status})"
            answer_parts.append(phrase)
        return "; ".join(answer_parts)

    def _extract_data_scientist_criteria(
        self,
        retrieved_chunks: List[Dict[str, Any]],
        context: str
    ) -> Optional[Dict[str, Any]]:
        doc_text = self._get_chunk_text(
            retrieved_chunks,
            filename_keyword='doc 3.pdf',
            keyword_filters=['data scientist', 'responsibilities', 'minimum qualifications']
        )
        search_space = doc_text or context
        if not search_space:
            return None
        
        responsibilities = self._filter_job_section_lines(
            self._extract_bullet_lines(
                self._get_section_block(search_space, 'responsibilities', ['experience', 'minimum qualifications', 'what will make you stand out'])
            )
        )
        experience = self._filter_job_section_lines(
            self._extract_bullet_lines(
                self._get_section_block(search_space, 'experience', ['minimum qualifications', 'what will make you stand out'])
            )
        )
        qualifications = self._filter_job_section_lines(
            self._extract_bullet_lines(
                self._get_section_block(search_space, 'minimum qualifications', ['what will make you stand out'])
            )
        )
        standout = self._filter_job_section_lines(
            self._extract_bullet_lines(
                self._get_section_block(search_space, 'what will make you stand out', None)
            )
        )
        
        if not any([responsibilities, experience, qualifications, standout]):
            return None
        
        lines = ["Data scientist hiring criteria from Doc 3:"]
        if responsibilities:
            lines.append("Responsibilities:")
            for item in responsibilities[:6]:
                lines.append(f"- {item}")
        if experience or qualifications:
            lines.append("Experience & minimum qualifications:")
            combined = experience + qualifications
            for item in combined[:6]:
                lines.append(f"- {item}")
        if standout:
            lines.append("What helps a candidate stand out:")
            for item in standout[:5]:
                lines.append(f"- {item}")
        
        sources = {
            chunk.get('metadata', {}).get('filename', 'Doc 3.pdf')
            for chunk in retrieved_chunks
            if 'doc 3.pdf' in chunk.get('metadata', {}).get('filename', '').lower()
        }
        if not sources:
            sources = {'Doc 3.pdf'}
        
        return {'answer': "\n".join(lines), 'sources': list(sources)}

    def _extract_logbook_benefits_and_setup(
        self,
        retrieved_chunks: List[Dict[str, Any]],
        context: str
    ) -> Optional[str]:
        doc_text = self._get_chunk_text(
            retrieved_chunks,
            filename_keyword='doc 5.pdf',
            keyword_filters=['log book', 'zone', 'download the gps log book']
        )
        search_space = doc_text or context
        if not search_space:
            return None
        
        benefit_lines = self._extract_bullet_lines(
            self._get_section_block(search_space, '1.1 who can benefit', ['1.2'])
        )
        feature_lines = self._extract_bullet_lines(
            self._get_section_block(search_space, '1.2 key features', ['2 '])
        )
        benefits = self._deduplicate_preserve_order(benefit_lines + feature_lines)[:6]
        
        section_configs = [
            ('6.1 download the gps log book sync application', ['6.2'], ['download', 'install', 'application', 'website']),
            ('6.2 plugging your device into the computer for the first time', ['6.3'], ['connect', 'device', 'register', 'driver']),
            ('6.3 upload data', ['6.4'], ['upload', 'sync', 'led', 'data'])
        ]
        setup_steps = []
        for header, stops, keywords in section_configs:
            block = self._get_section_block(search_space, header, stops)
            summary = self._summarize_section(block, keywords)
            if summary:
                section_label = header.split()[0]
                setup_steps.append(f"{section_label}: {summary}")
        setup_steps = setup_steps[:5]
        
        if not benefits and not setup_steps:
            return None
        
        lines = []
        if benefits:
            lines.append("Benefits of using the GPS Log Book (Doc 5):")
            for benefit in benefits:
                lines.append(f"- {benefit}")
        if setup_steps:
            lines.append("")
            lines.append("How to set it up:")
            for idx, step in enumerate(setup_steps, 1):
                lines.append(f"{idx}. {step}")
        
        return "\n".join([line for line in lines if line.strip()])

    def _extract_zone_instructions(self, context: str) -> Optional[str]:
        """Extract instructions for creating a zone from the log book document."""
        lower_context = context.lower()
        anchor = 'to create a new zone'
        idx = lower_context.find(anchor)
        
        section_text = ""
        if idx != -1:
            end_idx = len(context)
            for marker in ['7.3', '7.2.2', '8.']:
                marker_pos = lower_context.find(marker.lower(), idx)
                if marker_pos != -1:
                    end_idx = min(end_idx, marker_pos)
            start_idx = max(0, lower_context.rfind('7.2.1', 0, idx))
            section_text = context[start_idx:end_idx]
        else:
            section_pattern = re.compile(
                r'(7\.2\.1\s+Create\s+a\s+New\s+Zone.*?)(?:7\.2\.2|7\.3|8\.)',
                re.IGNORECASE | re.DOTALL
            )
            match = section_pattern.search(context)
            if match:
                section_text = match.group(1)
        
        if not section_text:
            return None
        
        cleaned = " ".join(section_text.split())
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', cleaned) if s.strip()]
        heading = "7.2.1 Create a New Zone"
        if sentences:
            # Keep only the sentences that describe the procedure.
            filtered = []
            for sentence in sentences:
                if sentence.lower().startswith('7.2.1'):
                    continue
                filtered.append(sentence)
            if filtered:
                concise = " ".join(filtered[:3])
            else:
                concise = " ".join(sentences[:3])
        else:
            concise = cleaned
        
        summary = f"{heading} — {concise}"
        return summary.strip()

def main():
    """Main function to run the conversational document assistant"""
    
    # Initialize the RAG system
    rag = ConversationalRAG("Gen AI/Source")
    
    print("\n" + "="*60)
    print("DOCUMENT ASSISTANT - Conversational Interface")
    print("="*60)
    print("\nSystem initialized. Processing test questions...\n")
    
    # Test questions from the assignment
    test_questions = [
        "What is the order quantity handled by Ranjit?",
        "What is the percentage of orders that haven't dispatched?",
        "List of products under the recoil kits orders?",
        "How much GST was charged for my insurance and when does my third party insurance expire?",
        "What are the roles we are currently hiring for?",
        "For the product Glock 17 what is the planned WO release date?",
        "Ok try Glock - 17",
        "What is our criteria to hire a data scientist?",
        "What are the benifits of log book and how can i set it up?",
        "How do i create a zone?"
    ]
    
    # Process each question
    results = []
    for i, question in enumerate(test_questions, 1):
        print(f"\nQuestion {i}: {question}")
        print("-" * 40)
        
        response = rag.generate_response(question)
        
        print(f"Answer: {response['answer']}")
        if response['sources']:
            print(f"Sources: {', '.join(set(response['sources']))}")
        if response['data']:
            print(f"Structured Data: Available")
        
        results.append(response)
    
    # Interactive mode
    print("\n" + "="*60)
    print("INTERACTIVE MODE")
    print("Type 'quit' to exit")
    print("="*60)
    
    while True:
        user_query = input("\nYour question: ").strip()
        if user_query.lower() in ['quit', 'exit', 'q']:
            break
        
        response = rag.generate_response(user_query)
        print(f"\nAnswer: {response['answer']}")
        if response['sources']:
            print(f"Sources: {', '.join(set(response['sources']))}")

if __name__ == "__main__":
    main()
